from fastapi import APIRouter, Depends, HTTPException, Query, Request, UploadFile, File
from app.database import get_db
from app.limiter import limiter
from app.middleware.auth import get_current_user
from app.middleware.validation import valid_object_id
from app.utils.errors import not_found, forbidden, bad_request, too_large, conflict
from app.models.schemas import (
    CVData,
    CVCreate,
    CVUpdate,
    CVHistoryRequest,
    AIPromptRequest,
    AIEditRequest,
    JobMatchRequest,
    CVReviewRequest,
    OptimizeBulletsRequest,
    KeywordGapRequest,
    AIInterviewRequest,
    CVAnalyticsCreate,
    CVAnalyticsEventOut,
    CVAnalyticsOut,
    CVKeywordTrendOut,
    SkillGapRequest,
    SkillGapResponse,
    SkillEndorsementCreate,
    SkillEndorsementOut,
    SectionSuggestionsResponse,
    SectionSuggestion,
    ToneAdjustRequest,
    ToneAdjustResponse,
)
from app.services import ai_service
from app.services.ats_service import simulateATS
from app.services.skill_market_service import fetch_market_data, find_courses
from datetime import datetime, timezone
from bson import ObjectId
import pdfplumber
import io
import re

MAGIC_PDF = b"%PDF"
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB

router = APIRouter()


def slugify(text: str) -> str:
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    return re.sub(r"[\s_-]+", "-", text)


def serialize_cv(cv: dict) -> dict:
    return {
        "id": str(cv["_id"]),
        "owner_id": cv["owner_id"],
        "owner_username": cv.get("owner_username", ""),
        "title": cv["title"],
        "data": cv["data"],
        "is_public": cv.get("is_public", False),
        "show_name": cv.get("show_name", True),
        "show_email": cv.get("show_email", False),
        "show_phone": cv.get("show_phone", False),
        "show_experience": cv.get("show_experience", True),
        "theme": cv.get("theme", "minimal"),
        "page_count": cv.get("page_count", 1),
        "template": cv.get("template", "standard"),
        "slug": cv.get("slug"),
        "created_at": cv["created_at"],
        "updated_at": cv["updated_at"],
    }


def serialize_analytics_event(doc: dict) -> CVAnalyticsEventOut:
    return CVAnalyticsEventOut(
        id=str(doc["_id"]),
        cv_id=doc["cv_id"],
        owner_id=doc["owner_id"],
        ats_score=doc.get("ats_score", 0),
        present_keywords=doc.get("present_keywords", []),
        missing_keywords=doc.get("missing_keywords", []),
        job_description=doc.get("job_description", ""),
        source=doc.get("source", "keyword_gap"),
        created_at=doc["created_at"],
    )


def keyword_trends(events: list[dict], field: str) -> list[CVKeywordTrendOut]:
    trends: dict[str, dict] = {}
    for event in events:
        created_at = event["created_at"]
        values = event.get(field, [])
        for value in values:
            if isinstance(value, str):
                keyword = value.strip()
                priority = "present"
                placement = ""
            else:
                keyword = str(value.get("keyword", "")).strip()
                priority = value.get("priority", "medium")
                placement = value.get("suggested_placement", "")
            if not keyword:
                continue
            key = keyword.lower()
            current = trends.setdefault(
                key,
                {
                    "keyword": keyword,
                    "count": 0,
                    "priority": priority,
                    "suggested_placement": placement,
                    "last_seen_at": created_at,
                },
            )
            current["count"] += 1
            if created_at > current["last_seen_at"]:
                current["keyword"] = keyword
                current["priority"] = priority
                current["suggested_placement"] = placement
                current["last_seen_at"] = created_at
    return [
        CVKeywordTrendOut(**item)
        for item in sorted(
            trends.values(),
            key=lambda item: (item["count"], item["last_seen_at"]),
            reverse=True,
        )
    ]


@router.post("")
async def create_cv(body: CVCreate, current_user=Depends(get_current_user), db=Depends(get_db)):
    now = datetime.now(timezone.utc)
    base_slug = slugify(body.title)
    slug = None
    if body.is_public:
        slug = f"{current_user['username']}-{base_slug}"

    cv_doc = {
        "owner_id": str(current_user["_id"]),
        "owner_username": current_user["username"],
        "title": body.title,
        "data": body.data.model_dump(),
        "is_public": body.is_public,
        "show_name": body.show_name,
        "show_email": body.show_email,
        "show_phone": body.show_phone,
        "show_experience": body.show_experience,
        "theme": body.theme,
        "page_count": body.page_count,
        "template": body.template,
        "slug": slug,
        "created_at": now,
        "updated_at": now,
    }
    result = await db.cvs.insert_one(cv_doc)
    cv_doc["_id"] = result.inserted_id
    return serialize_cv(cv_doc)


@router.get("")
async def list_cvs(
    skip: int = 0,
    limit: int = Query(20, ge=1, le=100),
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    user_id = str(current_user["_id"])
    total = await db.cvs.count_documents({"owner_id": user_id})
    cursor = db.cvs.find({"owner_id": user_id}).sort("updated_at", -1).skip(skip).limit(limit)
    cvs = await cursor.to_list(limit)
    return {
        "cvs": [serialize_cv(c) for c in cvs],
        "total": total,
        "skip": skip,
        "limit": limit,
    }


@router.get("/my-endorsements")
async def get_my_endorsements(current_user=Depends(get_current_user), db=Depends(get_db)):
    """Get all skills endorsed by the current user."""
    endorsements = await db.skill_endorsements.find(
        {"user_id": str(current_user["_id"])}
    ).sort("created_at", -1).limit(100).to_list(100)
    return [
        SkillEndorsementOut(
            id=str(e["_id"]),
            user_id=e["user_id"],
            skill=e["skill"],
            cv_id=e.get("cv_id"),
            endorser_username=e.get("endorser_username", ""),
            comment=e.get("comment", ""),
            created_at=e["created_at"],
        )
        for e in endorsements
    ]


@router.get("/{cv_id}")
async def get_cv(cv_id: str, current_user=Depends(get_current_user), db=Depends(get_db)):
    from bson import ObjectId
    if not ObjectId.is_valid(cv_id):
        raise HTTPException(status_code=400, detail="Invalid ID format")
    cv = await db.cvs.find_one({"_id": ObjectId(cv_id)})
    if not cv:
        raise not_found("CV")
    if cv["owner_id"] != str(current_user["_id"]) and current_user.get("role") not in ("admin", "superadmin"):
        if not cv.get("is_public"):
            raise forbidden("Access denied")
    return serialize_cv(cv)


@router.put("/{cv_id}")
async def update_cv(body: CVUpdate, cv_id: str, current_user=Depends(get_current_user), db=Depends(get_db)):
    from bson import ObjectId
    if not ObjectId.is_valid(cv_id):
        raise HTTPException(status_code=400, detail="Invalid ID format")
    cv = await db.cvs.find_one({"_id": ObjectId(cv_id)})
    if not cv:
        raise not_found("CV")
    if cv["owner_id"] != str(current_user["_id"]):
        raise forbidden("Access denied")

    # Save history snapshot before update (keep last 50 versions)
    await db.cv_history.insert_one({
        "cv_id": cv_id,
        "owner_id": str(current_user["_id"]),
        "snapshot": cv["data"],
        "title": cv["title"],
        "saved_at": datetime.now(timezone.utc),
    })
    # Enforce 50 version limit - delete oldest entries beyond limit
    count = await db.cv_history.count_documents({"cv_id": cv_id})
    if count > 50:
        oldest = await db.cv_history.find({"cv_id": cv_id}).sort("saved_at", 1).limit(count - 50).to_list(count - 50)
        if oldest:
            ids = [doc["_id"] for doc in oldest]
            await db.cv_history.delete_many({"_id": {"$in": ids}})

    updates = {"updated_at": datetime.now(timezone.utc)}
    if body.title is not None:
        updates["title"] = body.title
    if body.data is not None:
        updates["data"] = body.data.model_dump()
    if body.is_public is not None:
        updates["is_public"] = body.is_public
        if body.is_public:
            base_slug = slugify(body.title or cv["title"])
            updates["slug"] = f"{current_user['username']}-{base_slug}"
        else:
            updates["slug"] = None
    if body.show_name is not None:
        updates["show_name"] = body.show_name
    if body.show_email is not None:
        updates["show_email"] = body.show_email
    if body.show_phone is not None:
        updates["show_phone"] = body.show_phone
    if body.show_experience is not None:
        updates["show_experience"] = body.show_experience
    if body.theme is not None:
        updates["theme"] = body.theme
    if body.page_count is not None:
        updates["page_count"] = body.page_count
    if body.template is not None:
        updates["template"] = body.template    

    await db.cvs.update_one({"_id": ObjectId(cv_id)}, {"$set": updates})
    updated = await db.cvs.find_one({"_id": ObjectId(cv_id)})

    # Invalidate PDF cache since CV data/theme/template may have changed
    try:
        await db.pdf_cache.delete_many({"cv_id": cv_id})
    except Exception:
        pass

    return serialize_cv(updated)


@router.delete("/{cv_id}")
async def delete_cv(cv_id: str, current_user=Depends(get_current_user), db=Depends(get_db)):
    from bson import ObjectId
    if not ObjectId.is_valid(cv_id):
        raise HTTPException(status_code=400, detail="Invalid ID format")
    cv = await db.cvs.find_one({"_id": ObjectId(cv_id)})
    if not cv:
        raise not_found("CV")
    if cv["owner_id"] != str(current_user["_id"]) and current_user.get("role") not in ("admin", "superadmin"):
        raise forbidden("Access denied")
    await db.cv_history.delete_many({"cv_id": cv_id})
    await db.cv_analytics.delete_many({"cv_id": cv_id})
    await db.cvs.delete_one({"_id": ObjectId(cv_id)})
    return {"message": "CV deleted"}


@router.post("/{cv_id}/duplicate")
async def duplicate_cv(cv_id: str, current_user=Depends(get_current_user), db=Depends(get_db)):
    from bson import ObjectId
    if not ObjectId.is_valid(cv_id):
        raise HTTPException(status_code=400, detail="Invalid ID format")
    cv = await db.cvs.find_one({"_id": ObjectId(cv_id)})
    if not cv:
        raise not_found("CV")
    if cv["owner_id"] != str(current_user["_id"]):
        raise forbidden("Access denied")
    now = datetime.now(timezone.utc)
    new_doc = {
        "owner_id": str(current_user["_id"]),
        "owner_username": current_user["username"],
        "title": f"{cv['title']} (copy)",
        "data": cv["data"],
        "is_public": False,
        "theme": cv.get("theme", "minimal"),
        "page_count": cv.get("page_count", 1),
        "template": cv.get("template", "standard"), 
        "slug": None,
        "created_at": now,
        "updated_at": now,
    }
    result = await db.cvs.insert_one(new_doc)
    new_doc["_id"] = result.inserted_id
    return serialize_cv(new_doc)


@router.get("/{cv_id}/history")
async def get_cv_history(cv_id: str, current_user=Depends(get_current_user), db=Depends(get_db)):
    from bson import ObjectId
    if not ObjectId.is_valid(cv_id):
        raise HTTPException(status_code=400, detail="Invalid ID format")
    cv = await db.cvs.find_one({"_id": ObjectId(cv_id)})
    if not cv or cv["owner_id"] != str(current_user["_id"]):
        raise forbidden("Access denied")
    cursor = db.cv_history.find({"cv_id": cv_id}).sort("saved_at", -1).limit(20)
    history = await cursor.to_list(20)
    return [{"id": str(h["_id"]), "title": h.get("title"), "saved_at": h["saved_at"], "snapshot": h["snapshot"]} for h in history]


@router.post("/{cv_id}/history/restore")
async def restore_cv_version(
    body: CVHistoryRequest,
    cv_id: str,
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    """Restore CV to a previous version from history."""
    from bson import ObjectId
    if not ObjectId.is_valid(cv_id):
        raise HTTPException(status_code=400, detail="Invalid ID format")
    cv = await db.cvs.find_one({"_id": ObjectId(cv_id)})
    if not cv or cv["owner_id"] != str(current_user["_id"]):
        raise forbidden("Access denied")

    # Get the history snapshot
    history_doc = await db.cv_history.find_one({"_id": ObjectId(body.history_id), "cv_id": cv_id})
    if not history_doc:
        raise not_found("History entry")

    snapshot = history_doc.get("snapshot")
    if not snapshot:
        raise bad_request("Invalid history snapshot")

    # Save current state to history before restoring
    await db.cv_history.insert_one({
        "cv_id": cv_id,
        "owner_id": str(current_user["_id"]),
        "snapshot": cv["data"],
        "title": cv["title"],
        "saved_at": datetime.now(timezone.utc),
    })

    # Apply the snapshot
    await db.cvs.update_one(
        {"_id": ObjectId(cv_id)},
        {"$set": {"data": snapshot, "updated_at": datetime.now(timezone.utc)}},
    )

    return {"restored": True, "title": history_doc.get("title")}


@router.get("/{cv_id}/analytics", response_model=CVAnalyticsOut)
async def get_cv_analytics(cv_id: str, current_user=Depends(get_current_user), db=Depends(get_db)):
    from bson import ObjectId
    if not ObjectId.is_valid(cv_id):
        raise HTTPException(status_code=400, detail="Invalid ID format")
    cv = await db.cvs.find_one({"_id": ObjectId(cv_id)})
    if not cv or cv["owner_id"] != str(current_user["_id"]):
        raise forbidden("Access denied")
    events = await db.cv_analytics.find({"cv_id": cv_id, "owner_id": str(current_user["_id"])}).sort("created_at", -1).limit(50).to_list(50)
    return CVAnalyticsOut(
        cv_id=cv_id,
        events=[serialize_analytics_event(event) for event in events],
        missing_keyword_trends=keyword_trends(events, "missing_keywords"),
        present_keyword_trends=keyword_trends(events, "present_keywords"),
    )


@router.post("/{cv_id}/analytics", response_model=CVAnalyticsEventOut)
async def create_cv_analytics_event(body: CVAnalyticsCreate, cv_id: str, current_user=Depends(get_current_user), db=Depends(get_db)):
    from bson import ObjectId
    if not ObjectId.is_valid(cv_id):
        raise HTTPException(status_code=400, detail="Invalid ID format")
    cv = await db.cvs.find_one({"_id": ObjectId(cv_id)})
    if not cv or cv["owner_id"] != str(current_user["_id"]):
        raise forbidden("Access denied")
    now = datetime.now(timezone.utc)
    doc = {
        "cv_id": cv_id,
        "owner_id": str(current_user["_id"]),
        "ats_score": body.ats_score,
        "present_keywords": body.present_keywords,
        "missing_keywords": [item.model_dump() for item in body.missing_keywords],
        "job_description": body.job_description,
        "source": body.source,
        "created_at": now,
    }
    result = await db.cv_analytics.insert_one(doc)
    doc["_id"] = result.inserted_id
    return serialize_analytics_event(doc)


# AI Endpoints

@router.post("/ai/chat")
@limiter.limit("20/minute")
async def ai_chat(request: Request, body: AIPromptRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    result = await ai_service.chat_with_ai(
        body.message,
        cv_data=body.cv_data.model_dump() if body.cv_data else None,
        context=body.context or ""
    )
    usage = result.get("usage", {})
    tokens = usage.get("total_tokens", 0) or usage.get("completion_tokens", 0) or 0
    await db.ai_events.insert_one({
        "user_id": str(current_user["_id"]),
        "feature": "chat",
        "success": True,
        "tokens_approx": tokens,
        "ts": datetime.now(timezone.utc),
    })
    return {"response": result["response"], "usage": usage}


@router.post("/ai/chat/stream")
@limiter.limit("20/minute")
async def ai_chat_stream(request: Request, body: AIPromptRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    from fastapi.responses import StreamingResponse
    import asyncio
    
    async def event_stream():
        loop = asyncio.get_event_loop()
        gen = ai_service.chat_with_ai_stream(
            body.message,
            cv_data=body.cv_data.model_dump() if body.cv_data else None,
            context=body.context or "",
        )
        while True:
            try:
                event = await loop.run_in_executor(None, next, gen)
                yield event
            except StopIteration:
                break
    
    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/ai/generate-summary")
@limiter.limit("20/minute")
async def ai_generate_summary(request: Request, body: AIPromptRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    if not body.cv_data:
        raise bad_request("CV data required")
    summary = await ai_service.generate_summary(body.cv_data.model_dump())
    await db.ai_events.insert_one({
        "user_id": str(current_user["_id"]),
        "feature": "generate_summary",
        "success": True,
        "tokens_approx": 0,
        "ts": datetime.now(timezone.utc),
    })
    return {"summary": summary}


@router.post("/ai/edit")
@limiter.limit("20/minute")
async def ai_edit(request: Request, body: AIEditRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    updated = await ai_service.improve_cv_section(
        body.instruction, body.cv_data.model_dump(), body.section
    )
    await db.ai_events.insert_one({
        "user_id": str(current_user["_id"]),
        "feature": "edit",
        "success": True,
        "tokens_approx": 0,
        "ts": datetime.now(timezone.utc),
    })
    return {"data": updated}


@router.post("/ai/match-job")
@limiter.limit("20/minute")
async def ai_match_job(request: Request, body: JobMatchRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    updated = await ai_service.match_job_description(
        body.cv_data.model_dump(), body.job_description
    )
    await db.ai_events.insert_one({
        "user_id": str(current_user["_id"]),
        "feature": "match_job",
        "success": True,
        "tokens_approx": 0,
        "ts": datetime.now(timezone.utc),
    })
    return {"data": updated}


@router.post("/ai/interview")
@limiter.limit("20/minute")
async def ai_interview(request: Request, body: AIInterviewRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    response = await ai_service.interview_user(
        body.message,
        body.history,
        body.cv_data.model_dump() if body.cv_data else None,
    )
    await db.ai_events.insert_one({
        "user_id": str(current_user["_id"]),
        "feature": "interview",
        "success": True,
        "tokens_approx": 0,
        "ts": datetime.now(timezone.utc),
    })
    return {"response": response}


@router.post("/ai/review")
@limiter.limit("20/minute")
async def ai_review(request: Request, body: CVReviewRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    review = await ai_service.review_cv(
        body.cv_data.model_dump(),
        body.job_description or "",
    )
    await db.ai_events.insert_one({
        "user_id": str(current_user["_id"]),
        "feature": "review",
        "success": True,
        "tokens_approx": 0,
        "ts": datetime.now(timezone.utc),
    })
    return {"review": review}


@router.post("/ai/optimize-bullets")
@limiter.limit("20/minute")
async def ai_optimize_bullets(request: Request, body: OptimizeBulletsRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    updated = await ai_service.optimize_bullets(
        body.cv_data.model_dump(),
        body.experience_index,
    )
    await db.ai_events.insert_one({
        "user_id": str(current_user["_id"]),
        "feature": "optimize_bullets",
        "success": True,
        "tokens_approx": 0,
        "ts": datetime.now(timezone.utc),
    })
    return {"data": updated}


@router.post("/ai/keyword-gap")
@limiter.limit("20/minute")
async def ai_keyword_gap(request: Request, body: KeywordGapRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    analysis = await ai_service.keyword_gap_analysis(
        body.cv_data.model_dump(),
        body.job_description,
    )
    await db.ai_events.insert_one({
        "user_id": str(current_user["_id"]),
        "feature": "keyword_gap",
        "success": True,
        "tokens_approx": 0,
        "ts": datetime.now(timezone.utc),
    })
    return analysis


@router.post("/upload-cv")
async def upload_cv(file: UploadFile = File(...), current_user=Depends(get_current_user)):
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise too_large(f"File too large. Maximum size is {MAX_FILE_SIZE // 1024 // 1024}MB")

    filename = (file.filename or "").lower()

    # ── JSON import ────────────────────────────────────────────────────────
    if filename.endswith(".json"):
        try:
            data = json.loads(content.decode("utf-8"))
        except (json.JSONDecodeError, UnicodeDecodeError):
            raise bad_request("Invalid JSON file — could not parse")
        # Normalize into CVData shape
        try:
            validated = CVData(**data).model_dump()
        except Exception as e:
            raise bad_request(f"JSON does not match CV schema: {e}")
        return {"data": validated}

    # ── PDF import ─────────────────────────────────────────────────────────
    if content.startswith(MAGIC_PDF):
        text = ""
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            error_msg = str(e).lower()
            if "encrypt" in error_msg or "password" in error_msg:
                raise bad_request("Encrypted or password-protected PDFs are not supported")
            raise bad_request(f"Could not read PDF: {e}")
        if not text.strip():
            raise bad_request("Could not extract any text from PDF — it may be a scanned image")
        extracted = await ai_service.extract_cv_from_text(text)
        return {"data": extracted}

    # ── DOCX import ────────────────────────────────────────────────────────
    if filename.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_texts.append(" | ".join(cells))
            text = "\n".join(paragraphs + table_texts)
        except Exception as e:
            raise bad_request(f"Could not read DOCX file: {e}")
        if not text.strip():
            raise bad_request("Could not extract any text from DOCX")
        extracted = await ai_service.extract_cv_from_text(text)
        return {"data": extracted}

    raise bad_request("Unsupported file type. Accepted formats: PDF, DOCX, JSON")


# Skill Gap Engine

@router.post("/ai/skill-gap", response_model=SkillGapResponse)
@limiter.limit("20/minute")
async def analyze_skill_gaps(request: Request, body: SkillGapRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    # Fetch real job market data first (Adzuna, cached 1h), then run AI analysis with market context
    cv_data = body.cv_data.model_dump()
    target_role = body.target_role

    market_data = await fetch_market_data(target_role, db)
    analysis = await ai_service.analyze_skill_gaps(
        cv_data,
        target_role,
        market_data=market_data,
    )

    await db.ai_events.insert_one({
        "user_id": str(current_user["_id"]),
        "feature": "skill_gap",
        "success": True,
        "tokens_approx": 0,
        "ts": datetime.now(timezone.utc),
    })
    return SkillGapResponse(**analysis)


# ── Skill Endorsements ────────────────────────────────────────────────────────

@router.post("/ai/skill-endorsements", response_model=SkillEndorsementOut)
@limiter.limit("30/minute")
async def create_skill_endorsement(request: Request, body: SkillEndorsementCreate, current_user=Depends(get_current_user), db=Depends(get_db)):
    """Endorse a skill (self-endorsement or peer endorsement)."""
    now = datetime.now(timezone.utc)
    doc = {
        "user_id": str(current_user["_id"]),
        "skill": body.skill.strip().lower(),
        "cv_id": body.cv_id,
        "endorser_username": current_user["username"],
        "comment": body.comment.strip()[:200],
        "created_at": now,
    }
    # Upsert: one endorsement per user per skill
    await db.skill_endorsements.update_one(
        {"user_id": doc["user_id"], "skill": doc["skill"]},
        {"$set": doc},
        upsert=True,
    )
    created = await db.skill_endorsements.find_one(
        {"user_id": doc["user_id"], "skill": doc["skill"]}
    )
    return SkillEndorsementOut(
        id=str(created["_id"]),
        user_id=created["user_id"],
        skill=created["skill"],
        cv_id=created.get("cv_id"),
        endorser_username=created.get("endorser_username", ""),
        comment=created.get("comment", ""),
        created_at=created["created_at"],
    )


@router.get("/ai/skill-endorsements/{skill_name}")
@limiter.limit("30/minute")
async def get_skill_endorsements(request: Request, skill_name: str, current_user=Depends(get_current_user), db=Depends(get_db)):
    """Get endorsements for a specific skill across all users."""
    endorsements = await db.skill_endorsements.find(
        {"skill": skill_name.strip().lower()}
    ).sort("created_at", -1).limit(50).to_list(50)
    return {
        "skill": skill_name.strip().lower(),
        "count": len(endorsements),
        "endorsements": [
            SkillEndorsementOut(
                id=str(e["_id"]),
                user_id=e["user_id"],
                skill=e["skill"],
                cv_id=e.get("cv_id"),
                endorser_username=e.get("endorser_username", ""),
                comment=e.get("comment", ""),
                created_at=e["created_at"],
            )
            for e in endorsements
        ],
    }


@router.delete("/ai/skill-endorsements/{skill_name}")
@limiter.limit("30/minute")
async def remove_skill_endorsement(request: Request, skill_name: str, current_user=Depends(get_current_user), db=Depends(get_db)):
    """Remove a skill endorsement."""
    result = await db.skill_endorsements.delete_one({
        "user_id": str(current_user["_id"]),
        "skill": skill_name.strip().lower(),
    })
    return {"deleted": result.deleted_count > 0}


# ── Course suggestions ─────────────────────────────────────────────────────────

@router.get("/ai/courses/{skill_name}")
@limiter.limit("60/minute")
async def get_courses_for_skill(request: Request, skill_name: str):
    """Get learning course suggestions for a specific skill."""
    courses = find_courses(skill_name)
    return {"skill": skill_name.strip().lower(), "courses": courses}


# ATS Simulation

from pydantic import BaseModel


class ATSRequest(BaseModel):
    cv_data: dict
    job_description: str | None = None


class ATSFlagOut(BaseModel):
    severity: str
    category: str
    message: str
    details: str | None = None


class ATSVendorScoreOut(BaseModel):
    vendor: str
    score: int
    flags: list[ATSFlagOut] = []
    quirks_detected: list[str] = []


class ATSResultOut(BaseModel):
    score: int
    flags: list[ATSFlagOut]
    extracted_text: str
    section_headers_found: list[str]
    keyword_matches: list[dict]  # {keyword, tfidf, count}
    keyword_density: dict[str, float]
    missing_keywords: list[str]
    tfidf_scores: dict[str, float] = {}
    vendor_scores: list[ATSVendorScoreOut] = []
    vendor_overall: int = 0


@router.post("/ai/section-suggestions", response_model=SectionSuggestionsResponse)
@limiter.limit("20/minute")
async def ai_section_suggestions(request: Request, body: JobMatchRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    """Generate per-section suggestions to align CV to a job description."""
    suggestions = await ai_service.generate_section_suggestions(
        body.cv_data.model_dump(),
        body.job_description,
    )
    await db.ai_events.insert_one({
        "user_id": str(current_user["_id"]),
        "feature": "section_suggestions",
        "success": True,
        "tokens_approx": 0,
        "ts": datetime.now(timezone.utc),
    })
    return SectionSuggestionsResponse(suggestions=[SectionSuggestion(**s) for s in suggestions])


@router.post("/ai/adjust-tone", response_model=ToneAdjustResponse)
@limiter.limit("20/minute")
async def ai_adjust_tone(request: Request, body: ToneAdjustRequest, current_user=Depends(get_current_user), db=Depends(get_db)):
    """Adjust the tone/style of a specific CV section."""
    result = await ai_service.adjust_section_tone(
        body.cv_data.model_dump(),
        body.section,
        body.tone,
        body.custom_instruction or "",
    )
    await db.ai_events.insert_one({
        "user_id": str(current_user["_id"]),
        "feature": "adjust_tone",
        "success": True,
        "tokens_approx": 0,
        "ts": datetime.now(timezone.utc),
    })
    return ToneAdjustResponse(**result)


@router.post("/ai/ats-preview", response_model=ATSResultOut)
async def ats_preview(request: Request, body: ATSRequest, current_user=Depends(get_current_user)):
    """Simulate ATS parsing of a CV and return compatibility score with flagged issues."""
    # Parse CV data
    cv_data = CVData(**body.cv_data)

    # Run ATS simulation
    result = simulateATS(
        cv_data=cv_data,
        job_description=body.job_description,
    )

    # Convert to output model
    return ATSResultOut(
        score=result.score,
        flags=[
            ATSFlagOut(
                severity=f.severity,
                category=f.category,
                message=f.message,
                details=f.details,
            )
            for f in result.flags
        ],
        extracted_text=result.extracted_text,
        section_headers_found=result.section_headers_found,
        keyword_matches=result.keyword_matches,
        keyword_density=result.keyword_density,
        missing_keywords=result.missing_keywords,
        tfidf_scores=result.tfidf_scores,
        vendor_scores=[
            ATSVendorScoreOut(
                vendor=vs.vendor,
                score=vs.score,
                flags=[
                    ATSFlagOut(
                        severity=f.severity,
                        category=f.category,
                        message=f.message,
                        details=f.details,
                    )
                    for f in vs.flags
                ],
                quirks_detected=vs.quirks_detected,
            )
            for vs in result.vendor_scores
        ],
        vendor_overall=result.vendor_overall,
    )
