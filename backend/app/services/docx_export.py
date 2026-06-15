"""DOCX and plain text export service for CVs."""
import io
import re
from datetime import datetime, timezone
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.schemas import CVData


def strip_html(text: str) -> str:
    """Remove HTML tags from text for plain text export."""
    if not text:
        return ""
    # Remove HTML tags
    text = re.sub(r"<[^>]+>", "", text)
    # Decode common HTML entities
    text = text.replace("&nbsp;", " ")
    text = text.replace("&amp;", "&")
    text = text.replace("&lt;", "<")
    text = text.replace("&gt;", ">")
    text = text.replace("&quot;", '"')
    # Clean up whitespace
    text = re.sub(r"\s+", " ", text).strip()
    return text


def format_date_range(start: str, end: str, current: bool = False) -> str:
    """Format date range for display."""
    if current:
        return f"{start} - Present"
    if start and end:
        return f"{start} - {end}"
    if start:
        return start
    return end


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading to the document."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.size = Pt(14)
        run.font.bold = True


def _add_paragraph(doc: Document, text: str, bold: bool = False, size: int = 11) -> None:
    """Add a paragraph to the document."""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
    return para


def _add_bullet(doc: Document, text: str, indent: float = 0.5) -> None:
    """Add a bullet point to the document."""
    para = doc.add_paragraph(text, style="List Bullet")
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.size = Pt(11)
    # Set indentation
    para.paragraph_format.left_indent = Inches(indent)
    para.paragraph_format.first_line_indent = Inches(-0.25)


def generate_docx(
    cv_data: CVData,
    owner_username: str,
    cv_title: str,
    theme_name: str = "minimal",
    public_url: str = "",
) -> bytes:
    """Generate a DOCX file from CV data."""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Personal Info - Header
    personal = cv_data.personal_info
    name = strip_html(personal.full_name) or owner_username
    job_title = strip_html(personal.job_title)

    # Name as title
    title_para = doc.add_heading(name, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(18)
        run.font.bold = True

    if job_title:
        subtitle = doc.add_paragraph(job_title)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(14)

    # Contact info line
    contact_parts = []
    if personal.email:
        contact_parts.append(strip_html(personal.email))
    if personal.phone:
        contact_parts.append(strip_html(personal.phone))
    if personal.location:
        contact_parts.append(strip_html(personal.location))

    if contact_parts:
        contact_line = " | ".join(contact_parts)
        contact_para = doc.add_paragraph(contact_line)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(10)

    # Links
    links_parts = []
    if personal.linkedin:
        links_parts.append(f"LinkedIn: {strip_html(personal.linkedin)}")
    if personal.github:
        links_parts.append(f"GitHub: {strip_html(personal.github)}")
    if personal.portfolio:
        links_parts.append(f"Portfolio: {strip_html(personal.portfolio)}")
    if personal.website:
        links_parts.append(f"Website: {strip_html(personal.website)}")

    if links_parts:
        links_line = " | ".join(links_parts)
        links_para = doc.add_paragraph(links_line)
        links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in links_para.runs:
            run.font.size = Pt(9)

    doc.add_paragraph()  # Spacer

    # Summary section
    if cv_data.summary:
        _add_heading(doc, "Summary", level=1)
        summary_text = strip_html(cv_data.summary)
        _add_paragraph(doc, summary_text)
        doc.add_paragraph()

    # Experience section
    if cv_data.experience:
        _add_heading(doc, "Experience", level=1)
        for exp in cv_data.experience:
            # Company and dates
            company = strip_html(exp.company)
            role = strip_html(exp.role)
            dates = format_date_range(
                strip_html(exp.start_date), strip_html(exp.end_date), exp.current
            )

            # Role title line
            if role and company:
                header = f"{role} at {company}"
            elif role:
                header = role
            else:
                header = company

            if dates:
                header += f" | {dates}"

            _add_paragraph(doc, header, bold=True)

            # Description
            if exp.description:
                desc = strip_html(exp.description)
                _add_paragraph(doc, desc)

            # Achievements as bullets
            if exp.achievements:
                for achievement in exp.achievements:
                    ach_text = strip_html(achievement)
                    if ach_text:
                        _add_bullet(doc, ach_text)

            doc.add_paragraph()

    # Education section
    if cv_data.education:
        _add_heading(doc, "Education", level=1)
        for edu in cv_data.education:
            institution = strip_html(edu.institution)
            degree = strip_html(edu.degree)
            field = strip_html(edu.field)
            dates = format_date_range(
                strip_html(edu.start_date), strip_html(edu.end_date)
            )

            # Build education header
            if degree and field:
                header = f"{degree} in {field}"
            elif degree:
                header = degree
            else:
                header = institution

            if institution and degree:
                header = f"{institution}"
            elif institution:
                header = institution

            if degree or field:
                header += f" - {degree}" if degree else f" - {field}"

            if dates:
                header += f" | {dates}"

            _add_paragraph(doc, header, bold=True)

            if edu.grade:
                _add_paragraph(doc, f"Grade: {strip_html(edu.grade)}")

            if edu.description:
                desc = strip_html(edu.description)
                _add_paragraph(doc, desc)

            doc.add_paragraph()

    # Skills section
    if cv_data.skills:
        _add_heading(doc, "Skills", level=1)
        skills_text = ", ".join([strip_html(s) for s in cv_data.skills])
        _add_paragraph(doc, skills_text)
        doc.add_paragraph()

    # Certifications section
    if cv_data.certifications:
        _add_heading(doc, "Certifications", level=1)
        for cert in cv_data.certifications:
            name = strip_html(cert.name)
            issuer = strip_html(cert.issuer)
            date = strip_html(cert.date)

            if name:
                header = name
                if issuer:
                    header += f" - {issuer}"
                if date:
                    header += f" ({date})"
                _add_paragraph(doc, header, bold=True)

            if cert.credential_id:
                _add_paragraph(doc, f"Credential ID: {strip_html(cert.credential_id)}")

            doc.add_paragraph()

    # Projects section
    if cv_data.projects:
        _add_heading(doc, "Projects", level=1)
        for proj in cv_data.projects:
            name = strip_html(proj.name)
            if name:
                _add_paragraph(doc, name, bold=True)

            if proj.description:
                desc = strip_html(proj.description)
                _add_paragraph(doc, desc)

            if proj.technologies:
                techs = ", ".join([strip_html(t) for t in proj.technologies])
                _add_paragraph(doc, f"Technologies: {techs}")

            if proj.url:
                _add_paragraph(doc, f"URL: {strip_html(proj.url)}")

            doc.add_paragraph()

    # Awards section
    if cv_data.awards:
        _add_heading(doc, "Awards", level=1)
        for award in cv_data.awards:
            name = strip_html(award.name)
            issuer = strip_html(award.issuer)
            date = strip_html(award.date)

            if name:
                header = name
                if issuer:
                    header += f" - {issuer}"
                if date:
                    header += f" ({date})"
                _add_paragraph(doc, header, bold=True)

            if award.description:
                desc = strip_html(award.description)
                _add_paragraph(doc, desc)

            doc.add_paragraph()

    # Languages section
    if cv_data.languages:
        _add_heading(doc, "Languages", level=1)
        for lang in cv_data.languages:
            name = strip_html(lang.name)
            proficiency = strip_html(lang.proficiency)

            if name:
                if proficiency:
                    _add_paragraph(doc, f"{name} - {proficiency}")
                else:
                    _add_paragraph(doc, name)

        doc.add_paragraph()

    # Volunteer section
    if cv_data.volunteer:
        _add_heading(doc, "Volunteer Experience", level=1)
        for vol in cv_data.volunteer:
            org = strip_html(vol.organization)
            role = strip_html(vol.role)
            dates = format_date_range(
                strip_html(vol.start_date), strip_html(vol.end_date)
            )

            if role and org:
                header = f"{role} at {org}"
            elif role:
                header = role
            else:
                header = org

            if dates:
                header += f" | {dates}"

            if header:
                _add_paragraph(doc, header, bold=True)

            if vol.description:
                desc = strip_html(vol.description)
                _add_paragraph(doc, desc)

            doc.add_paragraph()

    # Add public URL at the end if available
    if public_url:
        doc.add_paragraph()
        url_para = doc.add_paragraph(f"View online: {public_url}")
        url_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in url_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 255)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_plain_text(
    cv_data: CVData,
    owner_username: str,
    cv_title: str,
) -> str:
    """Generate plain text from CV data (for ATS copy-paste)."""
    lines: list[str] = []

    # Personal Info
    personal = cv_data.personal_info
    name = strip_html(personal.full_name) or owner_username
    lines.append(name.upper())
    lines.append("")

    if personal.job_title:
        lines.append(strip_html(personal.job_title))
        lines.append("")

    # Contact info
    contact_parts = []
    if personal.email:
        contact_parts.append(strip_html(personal.email))
    if personal.phone:
        contact_parts.append(strip_html(personal.phone))
    if personal.location:
        contact_parts.append(strip_html(personal.location))

    if contact_parts:
        lines.append(" | ".join(contact_parts))
        lines.append("")

    # Links
    if personal.linkedin:
        lines.append(f"LinkedIn: {strip_html(personal.linkedin)}")
    if personal.github:
        lines.append(f"GitHub: {strip_html(personal.github)}")
    if personal.portfolio:
        lines.append(f"Portfolio: {strip_html(personal.portfolio)}")
    if personal.website:
        lines.append(f"Website: {strip_html(personal.website)}")

    if any([personal.linkedin, personal.github, personal.portfolio, personal.website]):
        lines.append("")

    # Summary
    if cv_data.summary:
        lines.append("--- SUMMARY ---")
        lines.append(strip_html(cv_data.summary))
        lines.append("")

    # Experience
    if cv_data.experience:
        lines.append("--- EXPERIENCE ---")
        for exp in cv_data.experience:
            company = strip_html(exp.company)
            role = strip_html(exp.role)
            dates = format_date_range(
                strip_html(exp.start_date), strip_html(exp.end_date), exp.current
            )

            if role and company:
                lines.append(f"{role} at {company} | {dates}")
            elif role:
                lines.append(f"{role} | {dates}")
            elif company:
                lines.append(f"{company} | {dates}")

            if exp.description:
                lines.append(strip_html(exp.description))

            if exp.achievements:
                for achievement in exp.achievements:
                    ach_text = strip_html(achievement)
                    if ach_text:
                        lines.append(f"  - {ach_text}")

            lines.append("")

    # Education
    if cv_data.education:
        lines.append("--- EDUCATION ---")
        for edu in cv_data.education:
            institution = strip_html(edu.institution)
            degree = strip_html(edu.degree)
            field = strip_html(edu.field)
            dates = format_date_range(
                strip_html(edu.start_date), strip_html(edu.end_date)
            )

            header_parts = []
            if institution:
                header_parts.append(institution)
            if degree:
                if field:
                    header_parts.append(f"{degree} in {field}")
                else:
                    header_parts.append(degree)

            header = " - ".join(header_parts)
            if dates:
                header += f" | {dates}"

            if header:
                lines.append(header)

            if edu.grade:
                lines.append(f"Grade: {strip_html(edu.grade)}")

            if edu.description:
                lines.append(strip_html(edu.description))

            lines.append("")

    # Skills
    if cv_data.skills:
        lines.append("--- SKILLS ---")
        skills_text = ", ".join([strip_html(s) for s in cv_data.skills])
        lines.append(skills_text)
        lines.append("")

    # Certifications
    if cv_data.certifications:
        lines.append("--- CERTIFICATIONS ---")
        for cert in cv_data.certifications:
            name = strip_html(cert.name)
            issuer = strip_html(cert.issuer)
            date = strip_html(cert.date)

            if name:
                parts = [name]
                if issuer:
                    parts.append(issuer)
                if date:
                    parts.append(f"({date})")
                lines.append(" - ".join(parts))

            if cert.credential_id:
                lines.append(f"Credential ID: {strip_html(cert.credential_id)}")

        lines.append("")

    # Projects
    if cv_data.projects:
        lines.append("--- PROJECTS ---")
        for proj in cv_data.projects:
            name = strip_html(proj.name)
            if name:
                lines.append(name)

            if proj.description:
                lines.append(strip_html(proj.description))

            if proj.technologies:
                techs = ", ".join([strip_html(t) for t in proj.technologies])
                lines.append(f"Technologies: {techs}")

            if proj.url:
                lines.append(f"URL: {strip_html(proj.url)}")

            lines.append("")

    # Awards
    if cv_data.awards:
        lines.append("--- AWARDS ---")
        for award in cv_data.awards:
            name = strip_html(award.name)
            issuer = strip_html(award.issuer)
            date = strip_html(award.date)

            if name:
                parts = [name]
                if issuer:
                    parts.append(issuer)
                if date:
                    parts.append(f"({date})")
                lines.append(" - ".join(parts))

            if award.description:
                lines.append(strip_html(award.description))

        lines.append("")

    # Languages
    if cv_data.languages:
        lines.append("--- LANGUAGES ---")
        for lang in cv_data.languages:
            name = strip_html(lang.name)
            proficiency = strip_html(lang.proficiency)

            if name:
                if proficiency:
                    lines.append(f"{name} - {proficiency}")
                else:
                    lines.append(name)

        lines.append("")

    # Volunteer
    if cv_data.volunteer:
        lines.append("--- VOLUNTEER EXPERIENCE ---")
        for vol in cv_data.volunteer:
            org = strip_html(vol.organization)
            role = strip_html(vol.role)
            dates = format_date_range(
                strip_html(vol.start_date), strip_html(vol.end_date)
            )

            if role and org:
                lines.append(f"{role} at {org} | {dates}")
            elif role:
                lines.append(f"{role} | {dates}")
            elif org:
                lines.append(f"{org} | {dates}")

            if vol.description:
                lines.append(strip_html(vol.description))

            lines.append("")

    # Join lines and clean up extra blank lines
    text = "\n".join(lines)
    # Remove trailing newlines
    text = text.rstrip("\n")
    return text